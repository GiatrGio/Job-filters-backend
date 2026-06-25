"""Extract plain text from an uploaded CV file (PDF / DOCX / TXT).

Runs entirely in memory: the bytes are parsed to text and the text is handed to
the CV-parse LLM call. Neither the file nor the extracted text is persisted —
only the structured, non-PII profile the LLM returns is stored (see
app/services/cv.py and migration 0014).
"""

from __future__ import annotations

import io

# Upload guardrails. The size cap bounds memory + parse cost; the text cap
# bounds the tokens we send to the LLM (a CV well over this is almost certainly
# padded and the head carries the signal that matters for fit).
MAX_CV_BYTES = 5 * 1024 * 1024  # 5 MB
MAX_CV_TEXT_CHARS = 60_000


class CvExtractionError(Exception):
    """Base class for recoverable CV upload problems (mapped to HTTP 400/415)."""


class UnsupportedCvFormat(CvExtractionError):
    pass


class CvFileTooLarge(CvExtractionError):
    pass


class EmptyCvText(CvExtractionError):
    pass


def _filename_ext(filename: str | None) -> str:
    if not filename or "." not in filename:
        return ""
    return filename.rsplit(".", 1)[1].lower()


def extract_cv_text(
    *,
    data: bytes,
    filename: str | None = None,
    content_type: str | None = None,
) -> str:
    """Return cleaned plain text from a CV upload.

    Raises CvFileTooLarge, UnsupportedCvFormat or EmptyCvText on bad input.
    """
    if len(data) > MAX_CV_BYTES:
        raise CvFileTooLarge(
            f"CV file is too large ({len(data)} bytes; max {MAX_CV_BYTES})."
        )
    if not data:
        raise EmptyCvText("The uploaded file is empty.")

    ext = _filename_ext(filename)
    ctype = (content_type or "").lower()

    if ext == "pdf" or "pdf" in ctype:
        text = _extract_pdf(data)
    elif ext == "docx" or "officedocument.wordprocessingml" in ctype:
        text = _extract_docx(data)
    elif ext in ("txt", "md", "text") or ctype.startswith("text/"):
        text = data.decode("utf-8", errors="replace")
    elif ext == "doc":
        # Legacy binary .doc isn't worth a parser dependency; ask for a re-export.
        raise UnsupportedCvFormat(
            "The old .doc format isn't supported. Please upload a PDF or .docx file."
        )
    else:
        raise UnsupportedCvFormat(
            "Unsupported file type. Please upload a PDF, .docx, or .txt file."
        )

    text = _normalize(text)
    if not text:
        # Almost always a scanned/image-only PDF with no text layer.
        raise EmptyCvText(
            "Couldn't read any text from this file. If it's a scanned PDF, "
            "please upload a text-based PDF or a .docx file."
        )
    return text[:MAX_CV_TEXT_CHARS]


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # corrupt / encrypted / unreadable
        raise UnsupportedCvFormat(f"Couldn't read this PDF ({exc}).") from exc


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise UnsupportedCvFormat(f"Couldn't read this .docx file ({exc}).") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _normalize(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()
